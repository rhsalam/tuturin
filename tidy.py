"""Rapikan transkrip tanpa memangkas isi.

Berbeda dari summarize.py: modul ini TIDAK meringkas. Setiap bagian sumber
harus muncul kembali di keluaran, hanya dibersihkan tanda baca, kata sisipan,
dan salah eja istilah. Karena keluaran sepanjang masukan, potongannya jauh
lebih kecil daripada peringkas agar muat di batas token balasan.

CATATAN PRIVASI: seperti summarize.py, modul ini mengirim transkrip ke DeepSeek.
"""
from __future__ import annotations

import re
from pathlib import Path

from summarize import SummaryError, _panggil

# Keluaran sepanjang masukan: ~1200 kata masuk menghasilkan ~1200 kata keluar,
# sekitar 2500 token balasan untuk Bahasa Indonesia.
CHUNK_WORDS = 1200
MAX_TOKENS = 4096

STEMPEL_TIAP = 60.0  # detik

SISTEM = (
    "Anda editor transkrip Bahasa Indonesia lisan. Tugas Anda merapikan, "
    "BUKAN meringkas. Setiap gagasan di sumber harus tetap ada di hasil. "
    "Anda tidak pernah menambah informasi yang tidak diucapkan."
)

ATURAN = """Rapikan transkrip berikut.

YANG DIPERBAIKI
- Tanda baca dan huruf besar.
- Buang kata sisipan yang hanya mengisi jeda: eh, anu, gitu, ya kan,
  apa namanya, nah, kan, ya. Ini WAJIB, bukan pilihan.
  Contoh:
    "Jadi gitu ya, kita itu nah harus usaha gitu." -> "Jadi kita harus usaha."
    "Apa namanya, eh, rezekinya itu kan dijamin gitu loh."
      -> "Rezekinya itu dijamin."
  Sisipan tetap dipertahankan HANYA bila membawa makna, misalnya "begitu
  caranya" atau "nah, ini penting" sebagai penanda pindah topik yang jelas.
- Buang pengulangan kata yang jelas kesalahan rekaman
  ("saya saya saya" -> "saya"). Pengulangan yang disengaja untuk penekanan
  tetap dipertahankan.
- Betulkan ejaan istilah: insya Allah, Subhanahu wa Ta'ala, Shallallahu
  'alaihi wa sallam, rezeki, omzet, cashflow, Alhamdulillah, Masya Allah,
  Assalamualaikum warahmatullahi wabarakatuh.

YANG DILARANG
- DILARANG meringkas, memotong, atau melewati bagian mana pun.
- DILARANG menambah kalimat, contoh, atau penjelasan yang tidak ada di sumber.
- DILARANG mengubah gaya bicara lisan menjadi bahasa formal atau baku.
  Kata khas lisan WAJIB dipertahankan apa adanya: banget, aja, enggak, nggak,
  kok, udah, gimana, kayak, dapet, sapaan "Bapak-Bapak sekalian".
  Membuang kata sisipan BUKAN memformalkan. Dua aturan ini tidak bertabrakan:
  buang pengisi jeda, pertahankan kosakata lisannya.
- Bagian yang tidak masuk akal atau jelas salah rekam: BIARKAN apa adanya
  lalu beri tanda [?] tepat sesudahnya. Jangan menebak maksudnya.

BENTUK KELUARAN
- Satu paragraf untuk setiap penanda waktu. Seluruh teks di antara dua
  penanda digabung menjadi SATU paragraf utuh.
- Setiap penanda waktu memulai paragraf baru. DILARANG meninggalkan penanda
  waktu di tengah paragraf.
- Jangan memecah lagi menjadi paragraf pendek.
- Bila topik berganti, sisipkan sub-judul Markdown "## Judul Topik" sebelum
  paragraf. Judul dirumuskan dari isi, bukan disalin mentah.
- Jangan menulis pengantar, catatan, atau komentar apa pun. Langsung isi.

Transkrip:
---
{teks}
---"""

POIN = """Berikut transkrip ceramah yang sudah dirapikan. Tulis {n} butir poin
utama dalam Markdown, satu baris per butir diawali "- ".

- Setiap butir satu kalimat padat yang berdiri sendiri.
- Hanya gunakan isi transkrip. Jangan menambah apa pun.
- Urutkan sesuai urutan pembahasan.
- Jangan menulis judul atau pengantar, langsung daftar butirnya.

Transkrip:
---
{teks}
---"""


def hitung_kata(t: str) -> int:
    """Hitung kata isi saja. Penanda waktu dikecualikan karena diulang di tiap
    paragraf pecahan, sehingga akan menggelembungkan angka "sesudah"."""
    return len(re.sub(r"\*{0,2}\[\d{2}:\d{2}:\d{2}\]\*{0,2}", " ", t).split())


def _stempel(detik: float) -> str:
    d = int(detik)
    return f"[{d // 3600:02d}:{(d % 3600) // 60:02d}:{d % 60:02d}]"


def sumber_berstempel(cues: list[dict]) -> tuple[str, bool]:
    """Susun teks sumber dengan penanda waktu tiap ~1 menit.

    Kembalikan (teks, ada_stempel). Bila cue tidak punya waktu yang masuk akal,
    teks tetap disusun tanpa penanda dan pemanggil diberi tahu.
    """
    cues = buang_ekor_berulang(cues)
    if not cues:
        return "", False
    punya_waktu = any(c.get("end", 0) > 0 for c in cues)
    bagian: list[str] = []
    berikutnya = 0.0
    for c in cues:
        teks = (c.get("text") or "").strip()
        if not teks:
            continue
        if punya_waktu and c.get("start", 0.0) >= berikutnya:
            bagian.append("\n" + _stempel(c["start"]))
            berikutnya = c["start"] + STEMPEL_TIAP
        bagian.append(teks)
    return " ".join(bagian).strip(), punya_waktu


def buang_ekor_berulang(cues: list[dict], min_ulang: int = 3) -> list[dict]:
    """Buang rentetan cue identik di ujung rekaman.

    Whisper kerap terjebak mengulang satu frasa pada keheningan atau musik
    penutup — pada satu ceramah 51 menit, "Sampai jumpa di video selanjutnya"
    muncul 14 kali berturut-turut di 16 detik terakhir. Pembersih satu-baris
    tidak menangkap pola ini.
    """
    if len(cues) < min_ulang:
        return cues
    akhir = (cues[-1].get("text") or "").strip().lower()
    if not akhir:
        return cues
    n = 0
    for c in reversed(cues):
        if (c.get("text") or "").strip().lower() == akhir:
            n += 1
        else:
            break
    return cues[:-n] if n >= min_ulang else cues


def potong(teks: str) -> list[str]:
    """Potong di batas penanda waktu agar tiap potongan punya stempel sendiri."""
    blok = re.split(r"(?=\[\d{2}:\d{2}:\d{2}\])", teks)
    blok = [b for b in blok if b.strip()]
    if not blok:
        return [teks] if teks.strip() else []
    bagian, kini, n = [], [], 0
    for b in blok:
        w = hitung_kata(b)
        if kini and n + w > CHUNK_WORDS:
            bagian.append(" ".join(kini)); kini, n = [], 0
        kini.append(b.strip()); n += w
    if kini:
        bagian.append(" ".join(kini))
    return bagian


def _bersihkan(hasil: str) -> str:
    """Buang pembungkus yang kadang ditambahkan model di luar instruksi."""
    hasil = re.sub(r"^```(?:markdown)?\s*|\s*```$", "", hasil.strip())
    hasil = re.sub(r"^(Berikut|Ini adalah|Hasil)[^\n]{0,80}:\s*\n+", "", hasil, count=1)
    return hasil.strip()


# Kata yang bila mendahului "gitu" membuatnya bermakna ("kayak gitu" = seperti
# itu), bukan sekadar pengisi jeda.
LINDUNGI = {"kayak", "kaya", "seperti", "macam", "begini", "begitu", "yang"}

# Sisipan sesuai daftar pengguna, hanya ketika menempel pada tanda baca —
# posisi khas pengisi jeda di akhir klausa.
SISIPAN_RE = re.compile(
    r"(\S+)(\s+)(gitu|ya kan|apa namanya|anu|eh)(\s*[.,!?;])", re.IGNORECASE)


def buang_sisipan(teks: str) -> tuple[str, int]:
    """Buang pengisi jeda di akhir klausa. Kembalikan (teks, jumlah dibuang).

    Sengaja konservatif: hanya menyentuh kemunculan sebelum tanda baca, dan
    melewati konteks yang membuat kata itu bermakna. Model bahasa terbukti
    terlalu segan membuang sisipan sendiri, sementara aturan ini bisa diuji.
    """
    n = 0

    def ganti(m: re.Match) -> str:
        nonlocal n
        sebelum = m.group(1).lower().strip('.,!?;:"\u201c\u201d')
        if sebelum in LINDUNGI:
            return m.group(0)
        n += 1
        return m.group(1) + m.group(4)

    # dua lintasan: menangkap rentetan seperti "gitu ya kan,"
    teks = SISIPAN_RE.sub(ganti, teks)
    teks = SISIPAN_RE.sub(ganti, teks)
    # Sisipan yang diapit koma meninggalkan tanda baca ganda.
    teks = re.sub(r"([,;])\s*[,;]+", r"\1", teks)
    teks = re.sub(r"[,;]\s*([.!?])", r"\1", teks)
    teks = re.sub(r"\s+([.,!?;])", r"\1", teks)
    return teks, n


STEMPEL_RE = re.compile(r"\[\d{2}:\d{2}:\d{2}\]")
# Akhir kalimat: tanda baca lalu spasi dan huruf besar / kurung siku penanda.
KALIMAT_RE = re.compile(r'(?<=[.!?])\s+(?=[A-Z"\u201c\[])')


def paragrafkan(isi: str) -> str:
    """Satu paragraf per penanda waktu, penanda ditebalkan dan tidak berulang.

    Model tidak konsisten menaati aturan bentuk, jadi penggabungan dikerjakan
    di sini. Sub-judul ditahan lalu ditempatkan tepat sebelum paragraf
    berikutnya, supaya tidak memotong satu menit menjadi dua paragraf yang
    memaksa penanda waktu diulang.
    """
    keluar: list[str] = []
    tertunda: list[str] = []
    buf: list[str] = []
    stempel = ""
    terakhir = ""

    def flush() -> None:
        nonlocal buf, terakhir
        teks = " ".join(buf).strip()
        buf = []
        if not teks:
            return
        keluar.extend(tertunda)
        tertunda.clear()
        if stempel and stempel != terakhir:
            keluar.append(f"**{stempel}** {teks}")
            terakhir = stempel
        else:
            keluar.append(teks)

    for blok in isi.split("\n\n"):
        b = blok.strip()
        if not b:
            continue
        if b.startswith("#"):
            flush()
            tertunda.append(b)
            continue
        for ruas in re.split(r"(?=\[\d{2}:\d{2}:\d{2}\])", b):
            r = ruas.strip()
            if not r:
                continue
            m = STEMPEL_RE.match(r)
            if m:
                # Penanda yang sama berarti masih menit yang sama: lanjutkan
                # paragraf berjalan, jangan memecahnya.
                if m.group(0) != stempel:
                    flush()
                    stempel = m.group(0)
                r = r[m.end():].strip()
            if r:
                buf.append(r)
    flush()
    keluar.extend(tertunda)
    return "\n\n".join(keluar)


def rapikan(cues: list[dict], key: str, lapor=None) -> tuple[str, dict]:
    teks, ada_stempel = sumber_berstempel(cues)
    if not teks:
        raise SummaryError("Transkrip kosong, tidak ada yang bisa dirapikan.")

    bagian = potong(teks)
    kata_awal = hitung_kata(teks)
    pakai = {"prompt_tokens": 0, "completion_tokens": 0}

    def catat(u):
        pakai["prompt_tokens"] += u.get("prompt_tokens", 0)
        pakai["completion_tokens"] += u.get("completion_tokens", 0)

    keluar: list[str] = []
    for i, b in enumerate(bagian, 1):
        if lapor:
            lapor(f"merapikan {i}/{len(bagian)}")
        h, u = _panggil(key, [
            {"role": "system", "content": SISTEM},
            {"role": "user", "content": ATURAN.format(teks=b)},
        ], suhu=0.2, max_tokens=MAX_TOKENS)
        catat(u)
        keluar.append(_bersihkan(h))

    isi, dibuang = buang_sisipan("\n\n".join(keluar))
    isi = paragrafkan(isi)

    if lapor:
        lapor("menyusun poin utama")
    # Poin dirumuskan dari hasil yang sudah rapi, bukan dari sumber mentah.
    ringkas_masukan = isi if hitung_kata(isi) <= 7000 else " ".join(isi.split()[:7000])
    poin, u = _panggil(key, [
        {"role": "system", "content": SISTEM},
        {"role": "user", "content": POIN.format(n=6, teks=ringkas_masukan)},
    ], suhu=0.3, max_tokens=900)
    catat(u)

    pakai.update({
        "fillers_removed": dibuang,
        "chunks": len(bagian),
        "words_before": kata_awal,
        "words_after": hitung_kata(isi),
        "has_timestamps": ada_stempel,
    })
    return _bersihkan(poin), (isi, pakai)


BAHASA = {"id": "Indonesia", "en": "Inggris", "ar": "Arab",
          "ja": "Jepang", "zh": "Mandarin"}


def _durasi_jam(teks_durasi: str, detik: float | None = None) -> str:
    if detik is None:
        m = re.match(r"(\d+)m\s*(\d+)d", teks_durasi or "")
        detik = int(m.group(1)) * 60 + int(m.group(2)) if m else 0
    d = int(detik)
    return f"{d // 3600:02d}:{(d % 3600) // 60:02d}:{d % 60:02d}"


def susun_markdown(judul: str, meta: dict, poin: str, isi: str) -> str:
    bahasa = BAHASA.get(meta["bahasa"], meta["bahasa"])
    jumlah = f"{meta['words_after']:,}".replace(",", ".")
    kepala = [
        f"# {judul} \u2014 Transkrip", "",
        f"**Sumber:** `{meta['filename']}` \u00b7 "
        f"**Durasi:** {_durasi_jam(meta.get('durasi',''), meta.get('detik'))} \u00b7 "
        f"**Bahasa:** {bahasa} \u00b7 **\u00b1 {jumlah} kata**", "",
        "> Transkrip otomatis (model Whisper, dijalankan offline) yang telah "
        "dirapikan. Istilah Arab, nama orang, dan angka sebaiknya dicek ulang "
        "terhadap audio aslinya.", "",
        "---", "", "## Poin Utama", "", poin, "", "---", "", isi, "",
        "---", "", f"*Dibuat dengan {MEREK}*", "",
    ]
    return "\n".join(kepala)

MEREK = "Tuturin"
LOGO = Path(__file__).parent / "static" / "favicon-192.png"


def _footer_merek(doc) -> None:
    """Footer halaman: lambang kecil + nama produk.

    Diletakkan di footer bagian, bukan di badan dokumen, supaya berulang di
    setiap halaman tanpa ikut tergeser saat isinya bertambah.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor, Inches

    par = doc.sections[0].footer.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO.exists():
        try:
            par.add_run().add_picture(str(LOGO), height=Inches(0.13))
            par.add_run("  ")
        except Exception:
            pass          # lambang hilang bukan alasan menggagalkan dokumen
    r = par.add_run(f"Dibuat dengan {MEREK}")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x8A, 0x8A, 0x8A)


def tulis_docx(path: Path, judul: str, meta: dict, poin: str, isi: str) -> None:
    """Tulis versi .docx yang bisa langsung disunting di Word."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    doc.add_heading(f"{judul} \u2014 Transkrip", level=0)

    tabel = doc.add_table(rows=0, cols=2)
    tabel.style = "Light List Accent 1"
    baris_meta = [
        ("Sumber", meta["filename"]),
        ("Durasi", _durasi_jam(meta.get("durasi", ""), meta.get("detik"))),
        ("Bahasa", BAHASA.get(meta["bahasa"], meta["bahasa"])),
        ("Jumlah kata", f"{meta['words_after']} (sumber {meta['words_before']})"),
        ("Penanda waktu",
         "ada, tiap ±1 menit" if meta["has_timestamps"] else "tidak ada di sumber"),
    ]
    for k, v in baris_meta:
        sel = tabel.add_row().cells
        sel[0].text = k
        sel[1].text = str(v)

    catatan = doc.add_paragraph()
    r = catatan.add_run(
        "Transkrip otomatis (model Whisper, dijalankan offline) yang telah "
        "dirapikan. Istilah Arab, nama orang, dan angka sebaiknya dicek ulang "
        "terhadap audio aslinya.")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_heading("Poin Utama", level=1)
    for baris in poin.splitlines():
        b = baris.strip().lstrip("-*").strip()
        if b:
            doc.add_paragraph(b, style="List Bullet")

    doc.add_page_break()

    for blok in isi.split("\n\n"):
        b = blok.strip()
        if not b:
            continue
        if b.startswith("#"):
            tingkat = len(b) - len(b.lstrip("#"))
            doc.add_heading(b.lstrip("#").strip(), level=min(max(tingkat, 1), 4))
            continue
        p = doc.add_paragraph()
        # Penanda waktu dibuat abu-abu agar mata langsung lompat ke isinya.
        m = re.match(r"^\*\*(\[\d{2}:\d{2}:\d{2}\])\*\*\s*(.*)$", b, re.S)
        if m:
            stempel = p.add_run(m.group(1) + " ")
            stempel.bold = True
            stempel.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            stempel.font.size = Pt(9)
            p.add_run(m.group(2))
        else:
            p.add_run(b)

    _footer_merek(doc)
    doc.save(str(path))


def _inline_ke_run(paragraf, teks: str) -> None:
    """Tulis teks ke paragraf, menghormati penanda **tebal** Markdown."""
    for i, potong in enumerate(re.split(r"\*\*(.+?)\*\*", teks)):
        if not potong:
            continue
        r = paragraf.add_run(potong)
        r.bold = i % 2 == 1      # potongan ganjil berasal dari dalam **...**


def tulis_ringkasan_docx(path: Path, judul: str, meta: dict, md: str) -> None:
    """Ringkasan AI -> .docx.

    Berbeda dari tulis_docx(): sumbernya Markdown bebas dari model, bukan
    struktur tetap. Yang dikenali: judul ##/###, daftar berpoin, dan **tebal**.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    doc.add_heading(f"{judul} — Ringkasan", level=0)

    tabel = doc.add_table(rows=0, cols=2)
    tabel.style = "Light List Accent 1"
    for k, v in (("Sumber", meta.get("filename", "-")),
                 ("Durasi", _durasi_jam("", meta.get("detik"))),
                 ("Bahasa", BAHASA.get(meta.get("bahasa", ""), meta.get("bahasa", "-")))):
        sel = tabel.add_row().cells
        sel[0].text = k
        sel[1].text = str(v)

    catatan = doc.add_paragraph()
    r = catatan.add_run(
        "Ringkasan dibuat mesin dari transkrip otomatis. Nama orang, angka, dan "
        "istilah sebaiknya diperiksa ulang terhadap audio aslinya.")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for baris in md.split("\n"):
        b = baris.strip()
        if not b or b == "---":
            continue
        j = re.match(r"^(#{1,4})\s+(.*)$", b)
        if j:
            doc.add_heading(j.group(2).strip(), level=min(len(j.group(1)), 4))
            continue
        p = re.match(r"^[-*]\s+(.*)$", b) or re.match(r"^\d+[.)]\s+(.*)$", b)
        if p:
            _inline_ke_run(doc.add_paragraph(style="List Bullet"), p.group(1))
            continue
        _inline_ke_run(doc.add_paragraph(), b)

    _footer_merek(doc)
    doc.save(str(path))
